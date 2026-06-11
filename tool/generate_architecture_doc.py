"""Generate VOX app architecture Word document."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "docs/VOX_App_Architecture.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("VOX Application — Architecture & API Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Overview of screens, external APIs, major flows (TTS, STT, Google Sign-In), "
        "and core patterns used across the Flutter app."
    )

    # --- App shell ---
    add_heading(doc, "1. App Shell (wraps every page)", 1)
    add_table(
        doc,
        ["Layer", "Role"],
        [
            ["Provider", "ThemeProvider, LanguageProvider, TtsService, TempLibraryProvider, TempNotesProvider, CustomCommandsProvider"],
            ["GlobalSttWrapper", "Assistant double-tap, mini player overlay, read-aloud voice listener"],
            ["MicRouteObserver", "Tracks current route for mic priority"],
            ["Firebase", "Auth state → load custom commands; logout → stop TTS"],
            ["app_links", "Magic-link email sign-in deep links"],
        ],
    )

    # --- Per page ---
    add_heading(doc, "2. APIs & Structures by Page", 1)
    add_table(
        doc,
        ["Page", "Route", "External APIs", "Firestore / Storage", "Key services & patterns"],
        [
            ["Splash", "home", "Firebase Auth", "—", "AppSession, AuthSession, AppBootstrap → Home or onboarding"],
            ["Home", "/home", "Firebase Auth, STT", "library, users/.../deleted_library", "StreamBuilder, FirestoreDataGate, TtsService, AppSpeechService, MicCoordinator, AnalyticsService"],
            ["Upload", "/upload", "file_picker, image_picker, ML Kit OCR, Syncfusion PDF", "library, users", "DocumentOcrService, DocumentTextExtractor, MicCoordinator, → ReaderPage"],
            ["Reader", "pushed", "flutter_tts via TtsService", "—", "RouteAware, suppresses mini player, LibraryHighlightService, AI sheet"],
            ["Notes", "/notes", "STT, record, just_audio, Groq Whisper", "notes, Storage recordings/{uid}/", "GroqService, TranscriptionQueue, PdfService, DocumentChatBuddySheet"],
            ["Dictionary", "/dictionary", "HTTP: Free Dictionary, Merriam-Webster, Wikipedia", "—", "AppSpeechService search mic, just_audio pronunciation"],
            ["Menu", "/menu", "Firebase Auth", "users", "Hub navigation, FloatingBotWrapper, logout → AuthSession"],
            ["Profile", "/profile", "Google Sign-In, image_picker", "users", "Guest upgrade, EmailChangeService, GuestUpgradeService"],
            ["Onboarding", "splash entry", "Google Sign-In, email/password, anonymous", "users", "UserProfilePage — first-run sign-in"],
            ["Contact Us", "/contact", "STT, flutter_tts, WhatsApp", "contact_messages → Cloud Function → EmailJS", "ContactInboxService → jamiremkanneh@gmail.com"],
            ["Recommendations", "/recommendations", "STT", "recommendations", "RecommendationsService (rating + message, no email)"],
            ["AI Result", "pushed", "OpenRouter + Groq fallback", "users/.../saved_docs", "AiService, SavedDocsService, local FlutterTts"],
            ["Saved Docs", "/saved_docs", "OpenRouter/Groq re-run", "users/.../saved_docs, assessments", "SavedDocsService, PdfService"],
            ["Statistics", "/statistics", "—", "users/.../analytics", "AnalyticsService (XP, streaks, achievements)"],
            ["Reminders", "/reminders", "flutter_local_notifications", "library, notes picker", "RemindersService, NotificationService"],
            ["Recycle Bin", "/recycle_bin", "—", "users/.../deleted_library", "FirestoreDataGate, restore to notes/library"],
            ["History", "/history", "—", "users/.../history", "Inline Firestore StreamBuilder"],
            ["Custom Commands", "/custom_commands", "—", "custom_commands", "CustomCommandsProvider, macro editor"],
            ["About", "/about", "—", "—", "Static UI + LanguageProvider"],
            ["FAQs", "/faqs", "OpenRouter/Groq chat", "—", "FloatingBotWrapper → AiService.askAssistant"],
            ["Mini player", "overlay", "—", "—", "MiniPlayerBar → globalNavigatorKey → ReaderPage"],
        ],
    )

    # --- External APIs ---
    add_heading(doc, "3. External API Summary", 1)
    add_table(
        doc,
        ["API", "Purpose"],
        [
            ["Firebase Auth", "Sign-in (Google, email, anonymous, magic link)"],
            ["Cloud Firestore", "Library, notes, users, devices, commands, contact, recommendations, analytics"],
            ["Firebase Storage", "Voice note recordings"],
            ["Google Sign-In", "OAuth for Firebase credential"],
            ["OpenRouter", "AI summary, flashcards, assistant, chat (primary)"],
            ["Groq", "AI fallback + Whisper transcription for voice notes"],
            ["Google ML Kit", "Camera/gallery OCR on Upload"],
            ["EmailJS", "Contact emails (server Cloud Function deliverContactMessage only)"],
            ["Dictionary HTTP", "freedictionaryapi.com, dictionaryapi.com, Wikipedia REST"],
            ["Sherpa-ONNX", "On-device keyword spotting during read-aloud"],
            ["speech_to_text", "All microphones via single AppSpeechService"],
            ["flutter_tts", "Global read-aloud (TtsService) + local TTS on some pages"],
            ["url_launcher", "WhatsApp on Contact Us"],
        ],
    )

    # --- Firestore collections ---
    add_heading(doc, "4. Firestore & Storage Collections", 1)
    add_table(
        doc,
        ["Collection / path", "Used by"],
        [
            ["users/{uid}", "Profile, menu, stats, onboarding"],
            ["users/{uid}/deleted_library", "Recycle bin, soft-delete"],
            ["users/{uid}/saved_docs", "Saved docs / AI outputs"],
            ["users/{uid}/analytics", "Statistics / gamification"],
            ["users/{uid}/history", "History screen"],
            ["devices/{deviceId}", "Device recognition / silent restore"],
            ["library", "Home, upload, reminders picker"],
            ["notes", "Notes, reminders picker"],
            ["custom_commands", "Custom voice commands"],
            ["contact_messages", "Contact Us → Cloud Function → EmailJS"],
            ["recommendations", "Play Store feedback ratings"],
            ["Storage: recordings/{uid}/", "Notes voice recordings"],
        ],
    )

    # --- TTS flow ---
    add_heading(doc, "5. TTS / Read-Aloud Flow", 1)
    doc.add_paragraph(
        "1. User taps Read on Home or Upload → TtsService.play()\n"
        "2. flutter_tts reads text in chunks; word/sentence highlights update via timers\n"
        "3. If user leaves reader: MiniPlayerBar shows on all tabs (via GlobalSttWrapper)\n"
        "4. If ReaderPage is open: setSuppressGlobalMiniPlayer(true) — full controls in reader\n"
        "5. Tap mini bar or ↑ button → ReaderPage (uses globalNavigatorKey)\n"
        "6. ReadingVoiceListener starts ReadAloudVoiceService (STT owner = reading)\n"
        "7. Keywords detected: pause, play, stop, forward, back, highlight\n"
        "8. ReadingVoiceController dispatches to TtsService (pause/resume/seek/stop)\n"
        "9. Custom voice commands via CommandDispatcher can also control TTS\n"
        "10. On speaker: TTS volume ducks while mic listens; restores after command"
    )
    doc.add_paragraph("Core types: TtsService, ReaderPage, MiniPlayerBar, ReadingVoiceListener, ReadAloudVoiceService, ReadingVoiceController, MicCoordinator.")

    # --- STT flow ---
    add_heading(doc, "6. STT / Microphone Flow", 1)
    doc.add_paragraph(
        "Single engine: AppSpeechService (wraps speech_to_text). Only one mic owner at a time.\n\n"
        "Mic owners:\n"
        "• assistant — GlobalSttWrapper (double-tap anywhere)\n"
        "• reading — ReadAloudVoiceService (hands-free while read-aloud)\n"
        "• home — Library search mic on /home\n"
        "• dictionary — Lookup mic on /dictionary\n"
        "• notes — Dictation + file recording (record package)\n"
        "• chatbot — Floating chat / Study Buddy sheet\n"
        "• contact_us / recommendations — Form field dictation\n\n"
        "MicCoordinator priority rules:\n"
        "• Read-aloud TTS playing → blocks assistant, search, notes, chatbot\n"
        "• Read-aloud paused → other mics may borrow mic\n"
        "• authFlowActive → blocks assistant during sign-in\n"
        "• externalCaptureActive → blocks all mics during camera scan\n\n"
        "Assistant flow: STT result → wake word strip → CommandDispatcher → custom commands / AiService.interpretVoiceAssistant → navigation or TTS macros"
    )

    # --- Google Sign-In ---
    add_heading(doc, "7. Google Sign-In Flow", 1)
    doc.add_paragraph(
        "1. Splash: AppSession.resolveLaunchDestination()\n"
        "2. If device recognized + auth restored → Home with welcome\n"
        "3. Else → UserProfilePage (onboarding)\n"
        "4. MicCoordinator.enterAuthFlow() — suspends other mics\n"
        "5. GoogleSignIn.authenticate() → idToken\n"
        "6. GoogleAuthProvider.credential → Firebase\n"
        "7. If anonymous user exists → linkWithCredential; else signInWithCredential\n"
        "8. _finishGoogleSignIn: write users/{uid} to Firestore\n"
        "9. AuthSession.markSignedIn() + AppSession.markSetupComplete()\n"
        "10. Link device in devices/{deviceId} → navigate Home\n\n"
        "Profile guest upgrade: same credential flow + GuestUpgradeService.migrateLocalData()\n"
        "Magic link: app_links → signInWithEmailLink → Home\n"
        "Logout: markPendingSignIn + signOut + enterAuthFlow\n\n"
        "Session layers:\n"
        "• AuthSession — SharedPreferences (guestMode, userId, hasProfile)\n"
        "• AppSession — devices/{deviceId} cold-start restore\n"
        "• AuthRestore — silent Google re-auth after splash (AppBootstrap.runDeferred)"
    )

    # --- Patterns ---
    add_heading(doc, "8. Major Patterns Used Everywhere", 1)
    add_table(
        doc,
        ["Pattern", "Where"],
        [
            ["Provider / ChangeNotifier", "Theme, language, TTS, temp caches, custom commands"],
            ["StreamBuilder", "Firestore lists (home, notes, recycle bin, reminders)"],
            ["FirestoreDataGate", "Permission/cache handling before showing lists"],
            ["Singleton services", "AppSpeechService, MicCoordinator, TtsService"],
            ["Route observers", "MicRouteObserver, appRouteObserver (Reader visibility)"],
            ["Guest vs signed-in", "AuthSession + TempLibraryProvider / TempNotesProvider"],
            ["MultiProvider push", "Reader opened with shared TtsService instance"],
        ],
    )

    # --- Navigation ---
    add_heading(doc, "9. Quick Navigation Map", 1)
    doc.add_paragraph(
        "Splash → Home (library hub)\n"
        "  ├─ FAB Upload → scan/OCR/PDF → Reader → TTS + voice commands\n"
        "  ├─ Notes → record → Groq transcribe → AI tools\n"
        "  ├─ Dictionary → HTTP APIs + search mic\n"
        "  ├─ Menu → Profile, Contact, Recommendations, Stats, Reminders, etc.\n"
        "  └─ Mini player (any tab) → expand ↑ → Reader"
    )

    doc.add_paragraph()
    p = doc.add_paragraph("Generated for Vox Application (voxapplication).")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    import os
    os.makedirs("docs", exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
